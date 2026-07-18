import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel,Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key not found")
client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"

job_discription="""
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.
Key job responsibilities
• Collaborate and communicate effectively with experienced cross-disciplinary Amazonians to design, build, and operate innovative products and services that delight our customers, while participating in technical discussions to drive solutions forward.
• Design and develop scalable solutions using cloud-native architectures and microservices in a large distributed computing environment.
• Participate in code reviews and contribute to technical documentation.
• Build and maintain resilient distributed systems that are scalable, fault-tolerant, and cost-effective.
• Leverage and contribute to the development of GenAI and AI-powered tools to enhance development productivity while staying current with emerging technologies.
• Write clean, maintainable code following best practices and design patterns.
• Work in an agile environment practicing CI/CD principles while participating in operational responsibilities including on-call duties.
• Demonstrate operational excellence through monitoring, troubleshooting, and resolving production issues.
Basic Qualifications
- Experience with at least one general-purpose programming language such as Java, Python, C++, C#, Go, Rust, or TypeScript
- Experience with data structure implementation, basic algorithm development, and/or object-oriented design principles
- Currently has, or is in the process of obtaining a bachelor’s degree in Computer Science, Computer Engineering, Data Science, Information Systems, or related STEM fields
- Must be 18 years of age of older
Preferred Qualifications
- Experience from previous technical internship(s) or demonstrated project experience
- Experience with one or more of the following: AI tools for development productivity, Cloud platforms (preferably AWS), Database systems (SQL and NoSQL), Contributing to open-source projects, Version control systems, Debugging and troubleshooting complex systems
- Demonstrated ability to learn and adapt to new technologies quickly
- Basic understanding of software development lifecycle (SDLC)
- Strong problem-solving and analytical skills
- Excellent written and verbal communication skills
"""

class jobD(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skills:list[str]
    minimum_experience:float | None
    education_requirements:list[str]
    responsibilities:list[str]

jobD_schema=jobD.model_json_schema()

system_prompt=f"""
You are an expert HR assistant.
Your job is to anyalyze the job discription and extract structured information from them.
Return only valid JSON schema:

{jobD_schema}

IMPORTANT:
Do not return the schema itself.
Do not return fields like "properties","title" or "type".

Fill the schema with the actual inforation extracted from the job discription.

If minimum experience is not mentioned return null.
If information for a particular list is not present, then just return an empty list.
Do not invent information.

"""
user_prompt=f"""
Analyze the following job discription:

{job_discription}

"""

message_system={
    "role":"system",
    "content":system_prompt
}

message_user={
    "role":"user",
    "content":user_prompt
}
response_format={
    "type":"json_object"
}

messages=[message_system,message_user]

response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)

import json
raw_json = response.choices[0].message.content
job_data=json.loads(raw_json)
job=jobD(**job_data)
print(job.minimum_experience)
print(job.education_requirements)

class MatchResults(BaseModel):
    score:float
    details:dict
class Experience(BaseModel):
    company:str | None=None
    role:str | None=None
    duration:str | None=None
    description:str | None=None
    skills_used:list[str]=[]

class Resume(BaseModel):
    name:str | None=None
    email:str | None=None
    phone:str | None=None

    total_experience_years:float | None=None

    skills:list[str]=[]
    experiences:list[Experience]=[]
    education:list[str]=[]
    projects:list[str]=[]
    certifications:list[str]=[]

resume_schema=Resume.model_json_schema()

def final_score(job,resume):
    match_schema=MatchResults.model_json_schema()
    prompt=f"""
    You are an expert HR recruiter.
    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}

    Return JSON matching this schema:
    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met or not
    5. Overall match percentage from 0 to 100
    6. A short final verdict about the resume
    7. Give some pointers of what all things can be improved

    Keep the response concise and easy to read
    """
    message={
        "role":"user",
         "content":prompt
    }
    messages=[message]
    response_format={
        "type":"json_object"

    }
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
    data=json.loads(response.choices[0].message.content)
    return MatchResults(**data)
def parse_resume(resume_text):
    system_prompt=f"""
    You are an expert resume parser.Extract information from the resume based on its meaning,not only based on exact section headings.
    For example:
    Experience
    Professional Experience
    Projects
    Internships
    Work History
    Employment

    These may all contain relevant experience

    Skills may appear in skills section , work experience section, internships,projects.

    Return only valid JSON format answers matching this schema:
    {resume_schema}


    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt=f"""
    Parse the following resume
    {resume_text}

    """

    message_system={
        "role":"system",
        "content":system_prompt
    }

    message_user={
        "role":"user",
        "content":user_prompt
    }
    messages=[message_system,message_user]
    response_format={
        "type":"json_object"
    }
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)

    raw_output=response.choices[0].message.content
    data=json.loads(raw_output)
    resume=Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader=PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text=page.extract_text()
        if page_text:
            text+=page_text+"\n"
    return text

def read_docx(file_path):
    document=Document(file_path)
    text=""
    for paragraph in document.paragraphs:
        if(paragraph.text.strip()):
            text+=paragraph.text+"\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower()==".pdf":
        return read_pdf(file_path)
    elif(file_path.suffix.lower()==".docx"):
        return read_docx(file_path)
    else:
        return None


resume_folder=Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf",".docx"]:
        continue
    print("\nPreproecessing:",file_path.name)
    resume_text=read_resume(file_path)
    parsed_resume=parse_resume(resume_text)#here i am making the first llm call
    time.sleep(30)
    result=final_score(job,parsed_resume)#here i am making the second llm call
    #now here i will implement a limimt to prevent the server of gpt to get jammed by millions of requests
    time.sleep(30)

    print("Score:",result.score)
    all_results.append({
        "name":parsed_resume.name,
        "score":result.score,
        "details":result.details
    })
    all_results.sort(
        key=lambda candidate:candidate["score"],
        reverse=True


    )
    top_2=all_results[:2]
    worst_2=all_results[-2:]

    print("TOP 2 CANDIDATES")
    for candidate in top_2:
        print(candidate["name"],"-",candidate["score"],"%")
        print(candidate["details"])

    print("LOWEST 2 CANDIDATES")
    for candidate in worst_2:
        print(candidate["name"],"-",candidate["score"],"%")
        print(candidate["details"])